"""
src/ingest.py
=============
Pipeline de ingestﺣ۲o de documentos jurﺣ­dicos para o RAG MROSC.

Fluxo:
  1. Carrega os documentos (DOCX e PDF) separadamente
  2. Parseia por artigo usando regex sensﺣ­vel ﺣ  estrutura legal brasileira
  3. Gera metadados por chunk: fonte, tipo, artigo, status (vigente/alterado/revogado)
  4. Aplica fallback de split por tamanho para artigos muito longos
  5. Gera embeddings e persiste no Chroma

Uso:
  python src/ingest.py
"""

import os
import re
import sys
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from tqdm import tqdm

# Carrega variﺣ۰veis de ambiente do .env (se existir)
load_dotenv()

# ---------------------------------------------------------------------------
# Configuraﺣ۶ﺣﭖes (lidas do .env ou valores padrﺣ۲o)
# ---------------------------------------------------------------------------
DATA_DIR = Path(__file__).parent.parent / "data"
LEI_FILE = DATA_DIR / "lei_mrosc.docx"
DECRETO_FILE = DATA_DIR / "decreto_57575_2016.pdf"

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))         # TODO(usuﺣ۰rio): ajuste se necessﺣ۰rio
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))    # TODO(usuﺣ۰rio): ajuste se necessﺣ۰rio

# ---------------------------------------------------------------------------
# Padrﺣﭖes de regex para textos jurﺣ­dicos brasileiros
# ---------------------------------------------------------------------------

# Detecta inﺣ­cio de artigo: "Art. 1ﺡﭦ", "Art. 22.", "Art. 35-A", "Artigo 10"
ARTICLE_PATTERN = re.compile(
    r"(?:^|\n)\s*"
    r"(Art(?:igo)?\.?\s*\d+[\w-]*\s*[ﺡﺍﺡﭦ]?\.?)"
    r"(?=\s)",
    re.MULTILINE | re.IGNORECASE,
)

# Detecta artigos revogados (ex: "Revogado pelo Decreto nﺡﭦ 61.615/2022")
REVOCATION_PATTERN = re.compile(
    r"(Revogad[ao]|Suprimid[ao])\s+(?:pelo|pela|por)\s+(?:Decreto|Lei|Portaria|Resoluﺣ۶ﺣ۲o)",
    re.IGNORECASE,
)

# Detecta artigos com redaﺣ۶ﺣ۲o alterada (ex: "Redaﺣ۶ﺣ۲o dada pelo Decreto nﺡﭦ 63.541/2024")
AMENDMENT_PATTERN = re.compile(
    r"(Redaﺣ۶ﺣ۲o\s+dada|Nova\s+redaﺣ۶ﺣ۲o|Alterado)\s+(?:pelo|pela|por)\s+(?:Decreto|Lei|Portaria)",
    re.IGNORECASE,
)

# Extrai nﺣﭦmero do artigo para metadado (ex: "Art. 22" ﻗ "22")
ARTICLE_NUMBER_PATTERN = re.compile(
    r"Art(?:igo)?\.?\s*(\d+[\w-]*)\s*[ﺡﺍﺡﭦ]?\.?",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Funﺣ۶ﺣﭖes de carregamento
# ---------------------------------------------------------------------------

def load_docx(file_path: Path) -> str:
    """Carrega texto de arquivo DOCX preservando parﺣ۰grafos."""
    try:
        import docx  # python-docx
    except ImportError:
        print("ﻗ Instale python-docx: pip install python-docx")
        sys.exit(1)

    doc = docx.Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def load_pdf(file_path: Path) -> str:
    """Carrega texto de arquivo PDF, pﺣ۰gina por pﺣ۰gina."""
    try:
        from pypdf import PdfReader
    except ImportError:
        print("ﻗ Instale pypdf: pip install pypdf")
        sys.exit(1)

    reader = PdfReader(str(file_path))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text.strip())
    return "\n".join(pages_text)


# ---------------------------------------------------------------------------
# Chunking sensﺣ­vel ﺣ  estrutura jurﺣ­dica
# ---------------------------------------------------------------------------

def detect_article_status(article_text: str) -> str:
    """
    Detecta o status de vigor de um artigo com base no prﺣﺏprio texto.
    Retorna: 'revogado', 'alterado' ou 'vigente'.
    """
    if REVOCATION_PATTERN.search(article_text):
        return "revogado"
    if AMENDMENT_PATTERN.search(article_text):
        return "alterado"
    return "vigente"


def extract_article_number(article_header: str) -> str:
    """Extrai o nﺣﭦmero do artigo do cabeﺣ۶alho detectado pelo regex."""
    match = ARTICLE_NUMBER_PATTERN.search(article_header)
    if match:
        return match.group(1)
    return "?"


def split_by_articles(
    full_text: str,
    fonte: str,
    tipo: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    """
    Divide o texto jurﺣ­dico em chunks por artigo.

    Para cada artigo encontrado:
    - Se couber no chunk_size: um ﺣﭦnico Document com metadados completos.
    - Se for muito longo: subdividido pelo RecursiveCharacterTextSplitter,
      mantendo os metadados do artigo em todos os sub-chunks.

    Retorna lista de Document com metadados: fonte, tipo, artigo, status.
    """
    # Localiza todas as ocorrﺣ۹ncias de "Art. X..." no texto
    matches = list(ARTICLE_PATTERN.finditer(full_text))

    documents: list[Document] = []

    # Texto antes do primeiro artigo (preﺣ۱mbulo/cabeﺣ۶alho)
    if matches:
        preamble = full_text[: matches[0].start()].strip()
        if preamble:
            documents.append(
                Document(
                    page_content=preamble,
                    metadata={
                        "fonte": fonte,
                        "tipo": tipo,
                        "artigo": "preﺣ۱mbulo",
                        "status": "vigente",
                        "secao": "preﺣ۱mbulo",
                    },
                )
            )

    fallback_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )

    for i, match in enumerate(matches):
        # Delimitaﺣ۶ﺣ۲o do artigo: do inﺣ­cio deste match atﺣ۸ o inﺣ­cio do prﺣﺏximo
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)

        article_text = full_text[start:end].strip()
        article_header = match.group(1)
        article_number = extract_article_number(article_header)
        status = detect_article_status(article_text)

        base_metadata = {
            "fonte": fonte,
            "tipo": tipo,
            "artigo": article_number,
            "artigo_header": article_header.strip(),
            "status": status,
        }

        # Se o artigo couber no limite de chunk, nﺣ۲o subdividir
        if len(article_text) <= chunk_size:
            documents.append(
                Document(page_content=article_text, metadata=base_metadata)
            )
        else:
            # Artigo muito longo: subdividir, preservando metadados
            sub_chunks = fallback_splitter.split_text(article_text)
            for j, sub_chunk in enumerate(sub_chunks):
                sub_metadata = {**base_metadata, "sub_chunk": j + 1}
                documents.append(
                    Document(page_content=sub_chunk, metadata=sub_metadata)
                )

    return documents


# ---------------------------------------------------------------------------
# Pipeline principal de ingestﺣ۲o
# ---------------------------------------------------------------------------

def ingest_documents(
    lei_path: Optional[Path] = None,
    decreto_path: Optional[Path] = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    """
    Carrega, parseia e gera chunks dos dois documentos jurﺣ­dicos.
    Retorna lista de Documents prontos para indexaﺣ۶ﺣ۲o.
    """
    lei_path = lei_path or LEI_FILE
    decreto_path = decreto_path or DECRETO_FILE

    all_documents: list[Document] = []

    # --- Lei Federal 13.019/2014 (DOCX) ---
    if not lei_path.exists():
        print(f"ﻗ ﺅﺕ  Arquivo da Lei nﺣ۲o encontrado: {lei_path}")
        print(
            "   Coloque 'lei_mrosc.docx' na pasta 'data/' antes de rodar a ingestﺣ۲o."
        )
    else:
        print(f"\nﻭ Carregando Lei 13.019/2014: {lei_path.name} ...")
        lei_text = load_docx(lei_path)
        lei_docs = split_by_articles(
            full_text=lei_text,
            fonte="Lei 13.019/2014",
            tipo="lei_federal",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        all_documents.extend(lei_docs)
        print(f"   ﻗ {len(lei_docs)} chunks gerados da Lei 13.019/2014")

        # Relatﺣﺏrio de status dos artigos
        revogados = [d for d in lei_docs if d.metadata["status"] == "revogado"]
        alterados = [d for d in lei_docs if d.metadata["status"] == "alterado"]
        if revogados:
            print(
                f"   ﻗ ﺅﺕ  Artigos detectados como revogados: "
                f"{', '.join(d.metadata['artigo'] for d in revogados)}"
            )
        if alterados:
            print(
                f"   ﻗﺗﺅﺕ  Artigos com redaﺣ۶ﺣ۲o alterada: "
                f"{', '.join(d.metadata['artigo'] for d in alterados)}"
            )

    # --- Decreto Municipal 57.575/2016 (PDF) ---
    if not decreto_path.exists():
        print(f"\nﻗ ﺅﺕ  Arquivo do Decreto nﺣ۲o encontrado: {decreto_path}")
        print(
            "   Coloque 'decreto_57575_2016.pdf' na pasta 'data/' antes de rodar a ingestﺣ۲o."
        )
    else:
        print(f"\nﻭ Carregando Decreto 57.575/2016: {decreto_path.name} ...")
        decreto_text = load_pdf(decreto_path)
        decreto_docs = split_by_articles(
            full_text=decreto_text,
            fonte="Decreto 57.575/2016",
            tipo="decreto_municipal",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        all_documents.extend(decreto_docs)
        print(f"   ﻗ {len(decreto_docs)} chunks gerados do Decreto 57.575/2016")

        revogados = [d for d in decreto_docs if d.metadata["status"] == "revogado"]
        alterados = [d for d in decreto_docs if d.metadata["status"] == "alterado"]
        if revogados:
            print(
                f"   ﻗ ﺅﺕ  Artigos detectados como revogados: "
                f"{', '.join(d.metadata['artigo'] for d in revogados)}"
            )
        if alterados:
            print(
                f"   ﻗﺗﺅﺕ  Artigos com redaﺣ۶ﺣ۲o alterada: "
                f"{', '.join(d.metadata['artigo'] for d in alterados)}"
            )

    return all_documents


def run_ingest_pipeline():
    """Executa o pipeline completo: ingestﺣ۲o + embeddings + Chroma."""
    from src.vectorstore import build_vectorstore  # importaﺣ۶ﺣ۲o local para evitar circular

    print("=" * 60)
    print("  RAG Jurﺣ­dico MROSC ﻗ Pipeline de Ingestﺣ۲o")
    print("=" * 60)

    docs = ingest_documents()

    if not docs:
        print(
            "\nﻗ Nenhum documento encontrado. "
            "Certifique-se de que os arquivos estﺣ۲o em 'data/'."
        )
        sys.exit(1)

    print(f"\nﻭ Total de chunks a indexar: {len(docs)}")
    print("\nﻭ۱ Gerando embeddings e indexando no Chroma ...")
    print(
        "   (O download do modelo de embeddings pode levar alguns minutos na 1ﺡ۹ execuﺣ۶ﺣ۲o)\n"
    )

    vectorstore = build_vectorstore(docs)

    print(f"\nﻗ Ingestﺣ۲o concluﺣ­da! Vector store salvo em: {os.getenv('CHROMA_PERSIST_DIR', './chroma_db')}")
    print("\nAgora vocﺣ۹ pode rodar:")
    print("  streamlit run app_streamlit.py    # interface principal")
    print("  python src/cli.py                 # debug via terminal")
    print("=" * 60)


if __name__ == "__main__":
    # Permite rodar como: python src/ingest.py
    # Adiciona o diretﺣﺏrio raiz ao path para imports absolutos
    sys.path.insert(0, str(Path(__file__).parent.parent))
    run_ingest_pipeline()
